"""DOCX file handler."""

import io
from typing import Any, Dict, List, Union

from docx import Document as DocxDocument

from app.models.enums import FileType
from app.utils.file_handlers.base import ExtractResult, FileHandler
from app.utils.logger import get_logger
from app.utils.validators import sanitize_text_content

logger = get_logger(__name__)


class DocxFileHandler(FileHandler):
    """Handler for DOCX files."""

    def can_handle(self, file_type: FileType) -> bool:
        return file_type == FileType.DOCX

    def _load_document(self, content: Union[str, bytes]) -> DocxDocument:
        if isinstance(content, str):
            content = content.encode("utf-8")
        return DocxDocument(io.BytesIO(content))

    def extract_text(self, content: Union[str, bytes], metadata: Dict[str, Any]) -> str:
        logger.debug("Extracting text from .docx file")
        doc = self._load_document(content)
        texts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        texts.append(cell_text)
        return sanitize_text_content("\n\n".join(texts))

    def extract_metadata(self, content: Union[str, bytes], filename: str) -> Dict[str, Any]:
        doc = self._load_document(content)
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return {
            "file_type": "docx",
            "filename": filename,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "char_count": len(text),
            "word_count": len(text.split()),
        }

    def extract_heading_map(self, content: Union[str, bytes]) -> List[Dict[str, Any]]:
        """Extract heading positions from a DOCX file for structural metadata indexing.

        Simulates the same paragraph-joining logic as extract_text() to produce
        character positions that match the stored document.content string.

        Returns a list of dicts with keys:
            pos   (int)  – character offset in the extracted text
            level (int)  – heading level 1-6
            text  (str)  – heading text
        sorted by position.
        """
        doc = self._load_document(content)
        return self._extract_headings_from_doc(doc)

    @staticmethod
    def _extract_headings_from_doc(doc: DocxDocument) -> List[Dict[str, Any]]:
        headings: List[Dict[str, Any]] = []
        current_pos = 0

        for p in doc.paragraphs:
            text = p.text
            if not text or not text.strip():
                continue

            style_name = p.style.name if p.style else ""
            if style_name.startswith("Heading "):
                try:
                    level = int(style_name.split(" ")[1])
                    if 1 <= level <= 6:
                        headings.append({"pos": current_pos, "level": level, "text": text.strip()})
                except (ValueError, IndexError) as exc:
                    logger.debug(
                        "Skipped DOCX heading with unparseable style '%s': %s", style_name, exc
                    )

            # Each kept paragraph contributes len(text) + 2 chars ("\n\n" separator)
            current_pos += len(text) + 2

        return headings

    def extract_all(self, content: Union[str, bytes], filename: str) -> ExtractResult:
        """One-pass DOCX extraction: load the document once, extract everything."""
        doc = self._load_document(content)

        texts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        texts.append(cell_text)
        text = sanitize_text_content("\n\n".join(texts))

        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        meta_text = "\n".join(p.text for p in doc.paragraphs if p.text)
        metadata = {
            "file_type": "docx",
            "filename": filename,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "char_count": len(meta_text),
            "word_count": len(meta_text.split()),
        }

        headings = self._extract_headings_from_doc(doc) or None
        return ExtractResult(text=text, metadata=metadata, headings=headings)
