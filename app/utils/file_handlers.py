"""File type handlers for document processing."""

import io
import statistics
from dataclasses import dataclass, field
from typing import Any, Dict, List, Optional, Union
from xml.etree import ElementTree

from docx import Document as DocxDocument

from app.models.enums import FileType
from app.utils.logger import get_logger
from app.utils.validators import sanitize_text_content

logger = get_logger(__name__)


class FileHandler:
    """Base class for file handlers."""

    def can_handle(self, file_type: FileType) -> bool:
        """Check if this handler can process the given file type."""
        raise NotImplementedError

    def extract_text(self, content: Union[str, bytes], metadata: Dict[str, Any]) -> str:
        """Extract text content from file."""
        raise NotImplementedError

    def extract_metadata(self, content: Union[str, bytes], filename: str) -> Dict[str, Any]:
        """Extract metadata from file."""
        raise NotImplementedError


class TextFileHandler(FileHandler):
    """Handler for plain text files (.txt)."""

    def can_handle(self, file_type: FileType) -> bool:
        """Check if this is a text file."""
        return file_type == FileType.TXT

    def extract_text(self, content: Union[str, bytes], metadata: Dict[str, Any]) -> str:
        """
        Extract text from plain text file.

        For .txt files, content is already text, just sanitize it.
        """
        logger.debug("Extracting text from .txt file")
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        return sanitize_text_content(content)

    def extract_metadata(self, content: Union[str, bytes], filename: str) -> Dict[str, Any]:
        """
        Extract metadata from text file.

        Returns basic metadata like line count, character count.
        """
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        lines = content.split("\n")
        return {
            "file_type": "txt",
            "filename": filename,
            "line_count": len(lines),
            "char_count": len(content),
            "word_count": len(content.split()),
        }


class MarkdownFileHandler(FileHandler):
    """Handler for Markdown files (.md)."""

    def can_handle(self, file_type: FileType) -> bool:
        """Check if this is a markdown file."""
        return file_type == FileType.MD

    def extract_text(self, content: Union[str, bytes], metadata: Dict[str, Any]) -> str:
        """
        Extract text from markdown file.

        For MVP: Keep markdown as-is for better context preservation.
        Future: Option to strip markdown syntax or convert to plain text.
        """
        logger.debug("Extracting text from .md file")
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        return sanitize_text_content(content)

    def extract_metadata(self, content: Union[str, bytes], filename: str) -> Dict[str, Any]:
        """
        Extract metadata from markdown file.

        Includes markdown-specific metadata like heading count.
        """
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        lines = content.split("\n")

        # Count headings
        heading_count = sum(1 for line in lines if line.strip().startswith("#"))

        # Check for frontmatter (YAML between --- markers)
        has_frontmatter = content.strip().startswith("---")

        return {
            "file_type": "md",
            "filename": filename,
            "line_count": len(lines),
            "char_count": len(content),
            "word_count": len(content.split()),
            "heading_count": heading_count,
            "has_frontmatter": has_frontmatter,
        }


class FB2FileHandler(FileHandler):
    """Handler for FB2 (FictionBook 2.0) files."""

    def can_handle(self, file_type: FileType) -> bool:
        return file_type == FileType.FB2

    @staticmethod
    def _parse_xml(content: str):
        """Parse FB2 XML with lxml recovery fallback.

        Tries stdlib ElementTree first (strict); if that fails, falls back to
        lxml with recover=True which silently fixes mismatched tags and similar
        structural errors common in downloaded FB2 files.
        """
        try:
            return ElementTree.fromstring(content)
        except ElementTree.ParseError:
            pass
        try:
            from lxml import etree

            root_lxml = etree.fromstring(
                content.encode("utf-8") if isinstance(content, str) else content,
                parser=etree.XMLParser(recover=True, encoding="utf-8"),
            )
            # Convert lxml tree to stdlib ElementTree via serialization
            reparsed = ElementTree.fromstring(etree.tostring(root_lxml))
            return reparsed
        except Exception:
            return None

    @staticmethod
    def _parse_body(body_node) -> List[Dict[str, Any]]:
        """Parse an FB2 <body> element into structured text segments.

        Recursively walks sections, extracting titles and paragraphs while
        tracking nesting depth for heading level attribution.

        Returns a list of dicts:
            text       (str)  – stripped text content
            is_heading (bool) – True for <title> elements inside <section>
            level      (int)  – 1-6 nesting depth (1 = top-level section)
        """
        segments: List[Dict[str, Any]] = []

        def _get_node_text(node) -> str:
            return " ".join(t.strip() for t in node.itertext() if t.strip())

        def _process(node, section_depth: int) -> None:
            local = node.tag.split("}")[-1] if "}" in node.tag else node.tag

            if local == "section":
                for child in node:
                    _process(child, section_depth + 1)
            elif local == "title":
                text = _get_node_text(node)
                if text:
                    segments.append(
                        {"text": text, "is_heading": True, "level": max(1, section_depth)}
                    )
            elif local in (
                "p",
                "v",
                "stanza",
                "subtitle",
                "text-author",
                "date",
                "epigraph",
                "cite",
                "annotation",
                "poem",
            ):
                text = _get_node_text(node)
                if text:
                    segments.append({"text": text, "is_heading": False, "level": 0})
            # image, binary, table, etc. — silently skipped (no meaningful text)

        for child in body_node:
            _process(child, 0)

        return segments

    def extract_text(self, content: Union[str, bytes], metadata: Dict[str, Any]) -> str:
        logger.debug("Extracting text from .fb2 file")
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        root = self._parse_xml(content)
        if root is None:
            return sanitize_text_content(content)

        parts: List[str] = []
        for node in root.iter():
            if node.tag.endswith("body"):
                parts.extend(seg["text"] for seg in self._parse_body(node))

        if not parts:
            # Fallback: dump all text from document
            parts.append(" ".join(t for t in root.itertext() if t.strip()))

        return sanitize_text_content("\n\n".join(parts))

    def extract_heading_map(self, content: Union[str, bytes]) -> List[Dict[str, Any]]:
        """Extract heading positions from an FB2 file for structural metadata indexing.

        Uses the same _parse_body traversal as extract_text() so that character
        positions match offsets in the stored document.content string.

        Returns a list of dicts with keys:
            pos   (int)  – character offset in the extracted text
            level (int)  – heading level 1-6 (section nesting depth)
            text  (str)  – heading text
        sorted by position.
        """
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        root = self._parse_xml(content)
        if root is None:
            return []

        all_segments: List[Dict[str, Any]] = []
        for node in root.iter():
            if node.tag.endswith("body"):
                all_segments.extend(self._parse_body(node))

        if not all_segments:
            return []

        headings: List[Dict[str, Any]] = []
        current_pos = 0
        for seg in all_segments:
            if seg["is_heading"] and 1 <= seg["level"] <= 6:
                headings.append({"pos": current_pos, "level": seg["level"], "text": seg["text"]})
            # Each segment contributes len(text) + 2 chars ("\n\n" separator)
            current_pos += len(seg["text"]) + 2

        return headings

    def extract_metadata(self, content: Union[str, bytes], filename: str) -> Dict[str, Any]:
        if isinstance(content, bytes):
            content = content.decode("utf-8", errors="replace")
        root = self._parse_xml(content)
        if root is None:
            return {
                "file_type": "fb2",
                "filename": filename,
                "char_count": len(content),
                "word_count": len(content.split()),
            }

        title = None
        author = None

        for node in root.iter():
            if node.tag.endswith("book-title") and node.text:
                title = node.text.strip()
                break

        for node in root.iter():
            if node.tag.endswith("author"):
                parts = []
                for child in node:
                    if child.tag.endswith("first-name") and child.text:
                        parts.append(child.text.strip())
                    if child.tag.endswith("last-name") and child.text:
                        parts.append(child.text.strip())
                if parts:
                    author = " ".join(parts)
                    break

        return {
            "file_type": "fb2",
            "filename": filename,
            "title": title,
            "author": author,
            "char_count": len(content),
            "word_count": len(content.split()),
        }


class DocxFileHandler(FileHandler):
    """Handler for DOCX files."""

    def can_handle(self, file_type: FileType) -> bool:
        return file_type == FileType.DOCX

    def _load_document(self, content: Union[str, bytes]) -> DocxDocument:
        if isinstance(content, str):
            content = content.encode("utf-8")
        return DocxDocument(io.BytesIO(content))

    def extract_text(self, content: Union[str, bytes], metadata: Dict[str, Any]) -> str:
        logger.debug("Extracting text from .docx file")
        doc = self._load_document(content)
        texts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        texts.append(cell_text)
        return sanitize_text_content("\n\n".join(texts))

    def extract_metadata(self, content: Union[str, bytes], filename: str) -> Dict[str, Any]:
        doc = self._load_document(content)
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        text = "\n".join(p.text for p in doc.paragraphs if p.text)
        return {
            "file_type": "docx",
            "filename": filename,
            "paragraph_count": paragraph_count,
            "table_count": table_count,
            "char_count": len(text),
            "word_count": len(text.split()),
        }

    def extract_heading_map(self, content: Union[str, bytes]) -> List[Dict[str, Any]]:
        """Extract heading positions from a DOCX file for structural metadata indexing.

        Simulates the same paragraph-joining logic as extract_text() to produce
        character positions that match the stored document.content string.

        Returns a list of dicts with keys:
            pos   (int)  – character offset in the extracted text
            level (int)  – heading level 1-6
            text  (str)  – heading text
        sorted by position.
        """
        doc = self._load_document(content)
        headings: List[Dict[str, Any]] = []
        current_pos = 0

        for p in doc.paragraphs:
            text = p.text
            if not text or not text.strip():
                continue

            style_name = p.style.name if p.style else ""
            if style_name.startswith("Heading "):
                try:
                    level = int(style_name.split(" ")[1])
                    if 1 <= level <= 6:
                        headings.append({"pos": current_pos, "level": level, "text": text.strip()})
                except (ValueError, IndexError):
                    pass

            # Each kept paragraph contributes len(text) + 2 chars ("\n\n" separator)
            current_pos += len(text) + 2

        return headings


@dataclass
class PDFExtractionProfile:
    """Per-document extraction parameters derived by _profile_document()."""

    # y0 threshold as fraction of page height — blocks starting above this are headers
    header_zone_fraction: float = 0.04
    # y0 threshold as fraction of page height — blocks starting here or below are footers
    footer_zone_fraction: float = 0.95
    # Normalised texts of confirmed running headers — excluded from text and headings
    running_header_texts: frozenset = field(default_factory=frozenset)
    # Font size ratio above which a block qualifies as a size-based heading
    size_ratio_threshold: float = 1.15
    # First-line bold fraction to qualify as a bold heading
    bold_fraction_threshold: float = 0.85
    # First-line length range [min, max) for bold headings
    bold_min_chars: int = 5
    bold_max_chars: int = 120
    # Heading texts appearing >= this many times kept only at first occurrence
    repeat_threshold: int = 4


class PDFFileHandler(FileHandler):
    """Handler for PDF files (.pdf).

    Uses PyMuPDF (fitz) directly for text extraction, with find_tables()
    to detect tables and render them as Markdown (preserving empty cells).
    Text blocks that overlap with detected table regions are skipped so
    they are not duplicated. Items are merged in vertical order.

    Heading detection uses font-size relative to the page median: blocks
    whose maximum span font size is noticeably larger than the median are
    treated as headings; level is estimated from the relative size.

    Scanned (image-only) PDFs are detected and rejected with a clear error.
    """

    def can_handle(self, file_type: FileType) -> bool:
        return file_type == FileType.PDF

    @staticmethod
    def _to_bytes(content: Union[str, bytes]) -> bytes:
        if isinstance(content, str):
            return content.encode("utf-8")
        return content

    @staticmethod
    def _normalize_block_text(block) -> str:
        """Return normalised single-line text for a PyMuPDF dict block.

        Concatenates span texts with no separator (matching PyMuPDF's own
        layout where word spacing is embedded inside span text), then collapses
        whitespace.  Used identically in _profile_document() and _extract_pdf()
        so that running-header lookup keys are guaranteed to match.
        """
        import re as _re

        raw = "".join(
            s.get("text", "") for ln in block.get("lines", []) for s in ln.get("spans", [])
        )
        return _re.sub(r"\s+", " ", raw).strip()

    @staticmethod
    def _profile_document(doc) -> "PDFExtractionProfile":
        """Scan document structure to derive per-document extraction parameters.

        Detects running headers/footers (text appearing on ≥30% of pages) and
        derives zone fractions that cleanly exclude them from extraction.
        """
        from collections import defaultdict

        import fitz

        page_count = len(doc)
        if page_count == 0:
            return PDFExtractionProfile()

        SCAN_ZONE = 0.10  # fraction of page height to scan for headers/footers
        REPEAT_FRAC = 0.30  # text on >= this fraction of pages → running element
        MARGIN_PT = 4.0  # extra clearance added to derived cutoff

        top_blocks: dict = defaultdict(list)  # normalised_text → [(y0, y1), …]
        bottom_blocks: dict = defaultdict(list)
        page_heights: list = []

        for page in doc:
            h = page.rect.height
            page_heights.append(h)
            top_cut = h * SCAN_ZONE
            bottom_cut = h * (1.0 - SCAN_ZONE)

            for b in page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE).get("blocks", []):
                if b.get("type") != 0:
                    continue
                y0, y1 = b["bbox"][1], b["bbox"][3]
                text = PDFFileHandler._normalize_block_text(b)
                if not text:
                    continue
                if y0 < top_cut:
                    top_blocks[text].append((y0, y1))
                if y1 > bottom_cut:
                    bottom_blocks[text].append((y0, y1))

        repeat_threshold = max(3, page_count * REPEAT_FRAC)
        median_h = statistics.median(page_heights) if page_heights else 800.0

        # Derive header zone from confirmed running headers
        running_header_texts: set = set()
        header_y1_max = 0.0
        for text, coords in top_blocks.items():
            if len(coords) >= repeat_threshold:
                running_header_texts.add(text)
                header_y1_max = max(header_y1_max, max(y1 for _, y1 in coords))

        header_zone_fraction = (
            min((header_y1_max + MARGIN_PT) / median_h, 0.10) if running_header_texts else 0.04
        )

        # Derive footer zone from confirmed running footers
        footer_y0_min = median_h
        for text, coords in bottom_blocks.items():
            if len(coords) >= repeat_threshold:
                footer_y0_min = min(footer_y0_min, min(y0 for y0, _ in coords))

        footer_zone_fraction = (
            max((footer_y0_min - MARGIN_PT) / median_h, 0.85) if footer_y0_min < median_h else 0.95
        )

        return PDFExtractionProfile(
            header_zone_fraction=header_zone_fraction,
            footer_zone_fraction=footer_zone_fraction,
            running_header_texts=frozenset(running_header_texts),
        )

    @staticmethod
    def _extract_footer_page_number(page) -> Optional[int]:
        """Try to extract the printed (logical) page number from a PDF page footer."""
        import re

        rect = page.rect
        footer_rect = type(rect)(0, rect.height * 0.95, rect.width, rect.height)
        text = page.get_text("text", clip=footer_rect).strip()
        if not text:
            return None
        for line in reversed(text.splitlines()):
            line = line.strip()
            if re.fullmatch(r"\d{1,4}", line):
                return int(line)
        return None

    @staticmethod
    def _rows_to_markdown(rows: List[List]) -> str:
        """Convert a list-of-lists table (from find_tables) to a Markdown table."""
        if not rows:
            return ""

        def cell(v: Any) -> str:
            return str(v).replace("|", "\\|").strip() if v is not None else ""

        header = [cell(c) for c in rows[0]]
        lines = [
            "| " + " | ".join(header) + " |",
            "| " + " | ".join("---" for _ in header) + " |",
        ]
        for row in rows[1:]:
            cells = [cell(c) for c in row]
            # Pad or trim to match header width
            while len(cells) < len(header):
                cells.append("")
            lines.append("| " + " | ".join(cells[: len(header)]) + " |")
        return "\n".join(lines)

    def _extract_pdf(self, content: bytes):
        """Extract (text, heading_map, page_map).

        For each page:
          1. Detect tables with find_tables(strategy="lines") → Markdown.
          2. Extract text blocks via get_text("dict"), skipping blocks that
             overlap ≥40% with a table region.
          3. Merge text blocks and table Markdowns sorted by top-y coordinate.
          4. Detect headings by comparing block font sizes to the page median.

        page_map: [[char_offset, physical_page, logical_page_or_null], ...]
        """
        import re

        import fitz

        try:
            doc = fitz.open(stream=content, filetype="pdf")
        except Exception as exc:
            raise ValueError(f"Cannot open PDF: {exc}") from exc

        has_text = any(page.get_text("text").strip() for page in doc)
        if not has_text:
            doc.close()
            raise ValueError(
                "No text could be extracted from this PDF. "
                "It may be a scanned document — OCR is not supported yet."
            )

        profile = self._profile_document(doc)
        logger.info(
            "PDF profile: header_zone=%.3f footer_zone=%.3f running_headers=%d",
            profile.header_zone_fraction,
            profile.footer_zone_fraction,
            len(profile.running_header_texts),
        )

        text_parts: List[str] = []
        headings: List[Dict[str, Any]] = []
        page_map: List[List] = []
        current_pos = 0

        for page in doc:
            page_num = page.number + 1
            page_start_pos = current_pos

            # ── 1. Detect tables ──────────────────────────────────────────
            table_rects: List[Any] = []  # fitz.Rect objects
            table_items: List = []  # (y0, markdown_string)
            try:
                tabs = page.find_tables(strategy="lines")
                for tab in tabs.tables:
                    rows = tab.extract()
                    if rows and len(rows) >= 2:
                        tr = fitz.Rect(tab.bbox)
                        table_rects.append(tr)
                        md = self._rows_to_markdown(rows)
                        table_items.append((tab.bbox[1], md))
            except Exception:
                pass

            # ── 2. Extract text blocks, skipping table regions ────────────
            text_items: List = []  # (y0, text)
            all_font_sizes: List[float] = []
            # (y0, block_text, max_font_size, first_line_text,
            #  first_line_bold_chars, first_line_total_chars)
            block_info: List = []

            page_h = page.rect.height
            header_cutoff = page_h * profile.header_zone_fraction
            footer_cutoff = page_h * profile.footer_zone_fraction

            try:
                block_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)
                for block in block_dict.get("blocks", []):
                    if block.get("type") != 0:
                        continue
                    brect = fitz.Rect(block["bbox"])

                    # Skip running-header and footer strips.
                    # Use y0 (block start) for the header zone so blocks that
                    # start inside the strip are filtered even if they extend
                    # slightly below the cutoff.
                    if brect.y0 < header_cutoff or brect.y0 >= footer_cutoff:
                        continue

                    # Skip blocks whose normalised text is a confirmed running header
                    if profile.running_header_texts:
                        if self._normalize_block_text(block) in profile.running_header_texts:
                            continue

                    # Skip if block overlaps substantially with a table
                    in_table = False
                    for trect in table_rects:
                        inter = brect & trect
                        if (
                            not inter.is_empty
                            and brect.get_area() > 0
                            and inter.get_area() / brect.get_area() >= 0.4
                        ):
                            in_table = True
                            break
                    if in_table:
                        continue

                    lines_text: List[str] = []
                    max_size = 0.0
                    first_line_text = ""
                    first_line_bold = 0
                    first_line_total = 0
                    for line in block.get("lines", []):
                        line_text = "".join(s.get("text", "") for s in line.get("spans", []))
                        is_first = not first_line_text and line_text.strip()
                        if line_text.strip():
                            lines_text.append(line_text.rstrip())
                            if is_first:
                                first_line_text = line_text.strip()
                        for span in line.get("spans", []):
                            sz = span.get("size", 0)
                            if sz > 0:
                                all_font_sizes.append(sz)
                                if sz > max_size:
                                    max_size = sz
                            if is_first:
                                n = len(span.get("text", ""))
                                first_line_total += n
                                if span.get("flags", 0) & 0b10000:  # bold bit
                                    first_line_bold += n

                    if lines_text:
                        block_text = "\n".join(lines_text)
                        block_info.append(
                            (
                                block["bbox"][1],
                                block_text,
                                max_size,
                                first_line_text,
                                first_line_bold,
                                first_line_total,
                            )
                        )
            except Exception:
                fallback = page.get_text("text")
                if fallback.strip():
                    text_items.append((0.0, fallback))

            # ── 3. Heading detection: font-size or bold first line ─────────
            median_size = statistics.median(all_font_sizes) if all_font_sizes else 0.0

            for y0, block_text, max_size, first_line_text, fl_bold, fl_total in block_info:
                text_items.append((y0, block_text))

                # Size-based: any font in block noticeably larger than median
                is_size_heading = (
                    median_size > 0 and max_size > median_size * profile.size_ratio_threshold
                )

                # Bold-based: first non-empty line of block is sufficiently bold
                # and short enough to be a heading (not a paragraph opening)
                fl_bold_fraction = fl_bold / fl_total if fl_total > 0 else 0.0
                is_bold_heading = (
                    fl_bold_fraction >= profile.bold_fraction_threshold
                    and profile.bold_min_chars <= len(first_line_text) < profile.bold_max_chars
                )

                if is_size_heading:
                    ratio = max_size / median_size
                    level = 1 if ratio >= 2.0 else (2 if ratio >= 1.5 else 3)
                    heading_text = re.sub(r"\s+", " ", block_text.strip())
                    if (
                        heading_text
                        and len(heading_text) < 200
                        and not re.fullmatch(r"\d+", heading_text)
                    ):
                        headings.append({"pos": current_pos, "level": level, "text": heading_text})
                elif is_bold_heading:
                    heading_text = re.sub(r"\s+", " ", first_line_text)
                    # Skip pure-number headings (page numbers, margin section
                    # numbers like "1", "2", "85" that weren't caught by the
                    # header/footer zone filter)
                    if heading_text and not re.fullmatch(r"\d+", heading_text):
                        headings.append({"pos": current_pos, "level": 2, "text": heading_text})

            # ── 4. Merge items and build page text ────────────────────────
            all_items = sorted(text_items + table_items, key=lambda x: x[0])
            page_text = "\n\n".join(item[1] for item in all_items).strip()

            if not page_text:
                continue

            logical_num = self._extract_footer_page_number(page)
            text_parts.append(page_text)
            page_map.append([page_start_pos, page_num, logical_num])
            current_pos += len(page_text) + 2  # +2 for \n\n page separator

        doc.close()

        # ── Deduplicate running headers ───────────────────────────────────
        # Heading texts that repeat on many pages are running headers (e.g.
        # "Unit 9", chapter titles in the margin).  Keep the first occurrence
        # and drop the rest so they don't pollute section_path in every chunk.
        repeat_threshold = profile.repeat_threshold
        text_counts: Dict[str, int] = {}
        for h in headings:
            text_counts[h["text"]] = text_counts.get(h["text"], 0) + 1
        seen: set = set()
        filtered: List[Dict[str, Any]] = []
        for h in headings:
            t = h["text"]
            if text_counts[t] >= repeat_threshold:
                if t not in seen:
                    seen.add(t)
                    filtered.append(h)  # keep only first occurrence
            else:
                filtered.append(h)
        headings = filtered

        full_text = sanitize_text_content("\n\n".join(text_parts))

        if len(full_text.strip()) < 100:
            raise ValueError(
                "Extracted text is too short (< 100 chars). "
                "This PDF may be a scanned document — OCR is not supported yet."
            )

        return full_text, headings, page_map

    def extract_text(self, content: Union[str, bytes], metadata: Dict[str, Any]) -> str:
        logger.debug("Extracting text from .pdf file")
        text, _, _ = self._extract_pdf(self._to_bytes(content))
        return text

    def extract_heading_map(self, content: Union[str, bytes]) -> List[Dict[str, Any]]:
        """Extract heading positions from a PDF for structural metadata indexing."""
        _, headings, _ = self._extract_pdf(self._to_bytes(content))
        return headings

    def extract_page_map(self, content: Union[str, bytes]) -> List[List[int]]:
        """Extract page boundary map: [[char_offset, page_number], ...] (1-indexed).

        Each entry marks the character offset in the extracted text where a new
        page begins.  Use _get_page_for_char() in document_processor to look up
        the page number for any chunk by its start_char.
        """
        _, _, page_map = self._extract_pdf(self._to_bytes(content))
        return page_map

    def extract_metadata(self, content: Union[str, bytes], filename: str) -> Dict[str, Any]:
        import fitz

        try:
            doc = fitz.open(stream=self._to_bytes(content), filetype="pdf")
        except Exception:
            return {"file_type": "pdf", "filename": filename}

        meta = doc.metadata or {}
        page_count = doc.page_count
        doc.close()

        return {
            "file_type": "pdf",
            "filename": filename,
            "page_count": page_count,
            "title": meta.get("title") or None,
            "author": meta.get("author") or None,
            "creator": meta.get("creator") or None,
        }


class FileHandlerFactory:
    """
    Factory for creating appropriate file handlers.

    Usage:
        handler = FileHandlerFactory.get_handler(FileType.MD)
        text = handler.extract_text(content, {})
    """

    _handlers = [
        TextFileHandler(),
        MarkdownFileHandler(),
        FB2FileHandler(),
        DocxFileHandler(),
        PDFFileHandler(),
    ]

    @classmethod
    def get_handler(cls, file_type: FileType) -> FileHandler:
        """
        Get appropriate handler for file type.

        Args:
            file_type: Type of file to handle

        Returns:
            FileHandler instance

        Raises:
            ValueError: If no handler found for file type
        """
        for handler in cls._handlers:
            if handler.can_handle(file_type):
                return handler

        raise ValueError(f"No handler found for file type: {file_type}")

    @classmethod
    def register_handler(cls, handler: FileHandler) -> None:
        """
        Register a new file handler.

        Useful for adding handlers for new file types in later phases.

        Args:
            handler: FileHandler instance to register
        """
        cls._handlers.append(handler)
        logger.info(f"Registered file handler: {handler.__class__.__name__}")


def process_file(content: Union[str, bytes], filename: str, file_type: FileType) -> Dict[str, Any]:
    """
    Process a file and extract text and metadata.

    Convenience function that uses FileHandlerFactory.

    Args:
        content: File content as string
        filename: Name of the file
        file_type: Type of the file

    Returns:
        Dictionary with 'text' and 'metadata' keys

    Example:
        >>> result = process_file(content, "test.md", FileType.MD)
        >>> print(result['text'])
        >>> print(result['metadata'])
    """
    handler = FileHandlerFactory.get_handler(file_type)

    # Extract metadata first
    metadata = handler.extract_metadata(content, filename)

    # Extract text
    text = handler.extract_text(content, metadata)

    return {
        "text": text,
        "metadata": metadata,
    }
